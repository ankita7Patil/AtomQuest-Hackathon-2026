from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


OUT = "AtomQuest_GoalTrack_Submission_Final.docx"


def set_cell(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(19, 33, 31)
    return paragraph


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AtomQuest Hackathon 2026 Submission")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(21)
run.font.color.rgb = RGBColor(19, 33, 31)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("GoalTrack Portal - Company Goal Setting & Tracking Platform")
run.font.name = "Arial"
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(53, 103, 91)

add_heading(doc, "Project Snapshot", 1)
doc.add_paragraph(
    "GoalTrack Portal is a full-stack company goal management system built for structured performance planning, manager approval, quarterly check-ins, and transparent progress tracking. "
    "It supports Admin, Manager, and Employee roles with JWT authentication, real user management, validated goal plans, planned vs actual tracking, audit logs, and CSV reports."
)

table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
rows = [
    ("Deployed frontend", "https://atom-quest-hackathon-2026.vercel.app"),
    ("Deployed backend", "https://atomquest-hackathon-2026-pgnl.onrender.com"),
    ("Backend health check", "https://atomquest-hackathon-2026-pgnl.onrender.com/api/health"),
    ("Source repository", "https://github.com/ankita7Patil/AtomQuest-Hackathon-2026"),
]
for row, (label, value) in zip(table.rows, rows):
    set_cell(row.cells[0], label, True)
    set_cell(row.cells[1], value)

add_heading(doc, "Evaluator Login Credentials", 1)
creds = doc.add_table(rows=4, cols=3)
creds.style = "Table Grid"
for i, text in enumerate(["Role", "Email", "Password"]):
    set_cell(creds.rows[0].cells[i], text, True)
for row, data in enumerate(
    [
        ("Admin", "admin@atomquest.com", "Admin@12345"),
        ("Manager", "manager2@test.com", "Manager@123"),
        ("Employee", "employee2@test.com", "Employee@123"),
    ],
    start=1,
):
    for col, text in enumerate(data):
        set_cell(creds.rows[row].cells[col], text)

add_heading(doc, "Implemented Features", 1)
features = [
    "Login authentication using JWT",
    "Role-based dashboards for Employee, Manager, and Admin",
    "Admin user management for creating real company users",
    "Goal creation, submission, approval, and rejection",
    "Validation rules: total weightage equals 100%, minimum goal weightage is 10%, maximum 8 goals",
    "Quarterly check-ins for Q1, Q2, Q3, and Q4",
    "Planned vs actual tracking with weighted progress calculation",
    "Status updates: Not Started, On Track, Completed",
    "Shared goals support",
    "Audit trail logging",
    "CSV/Excel-ready report export",
    "Responsive modern UI",
    "MongoDB Atlas schemas and deployment-ready REST API",
    "Production deployment on Vercel and Render",
]
for item in features:
    doc.add_paragraph(item, style="List Bullet")

add_heading(doc, "Architecture Diagram", 1)
arch = doc.add_table(rows=5, cols=3)
arch.style = "Table Grid"
arch.alignment = WD_TABLE_ALIGNMENT.CENTER
diagram = [
    ("Users", "React Portal on Vercel", "Express API on Render"),
    ("Admin", "User Management", "JWT + RBAC"),
    ("Employee", "Goals + Check-ins", "Validation Workflow"),
    ("Manager", "Approvals + Reports", "Audit Logging"),
    ("All Roles", "CSV Export", "MongoDB Atlas"),
]
for r, row_data in enumerate(diagram):
    for c, text in enumerate(row_data):
        set_cell(arch.rows[r].cells[c], text, bold=(r == 0 or c == 0))

doc.add_paragraph("Flow: User -> React + Tailwind Portal -> Express REST API -> JWT Authentication -> Role Protection -> User/Goal/Check-in/Report Services -> MongoDB Atlas.")

add_heading(doc, "Role-Based Workflow", 1)
workflow = [
    "Admin creates managers, employees, and other admins, then assigns employees to managers.",
    "Employee creates goals, follows weightage validation, and submits the plan for review.",
    "Manager reviews submitted goals and approves or rejects them.",
    "Employee updates quarterly planned vs actual progress and goal status.",
    "Admin and managers export CSV reports and review audit trail activity.",
]
for item in workflow:
    doc.add_paragraph(item, style="List Bullet")

add_heading(doc, "Deployment Plan", 1)
steps = [
    "Frontend deployed on Vercel from the client folder.",
    "Backend deployed on Render from the server folder.",
    "MongoDB Atlas used as the cloud database.",
    "VITE_API_URL points the frontend to the Render backend.",
    "Render environment variables include MONGODB_URI, JWT_SECRET, CLIENT_URL, and first-admin credentials.",
]
for item in steps:
    doc.add_paragraph(item, style="List Number")

doc.save(OUT)
print(OUT)
